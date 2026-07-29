# tests/test_audiobook_extraction.py
# ===================================
# Every extractor returns the same shape (ExtractionResult), so these tests
# all follow the same pattern: build a real source file in tmp_path, run the
# extractor, assert on titles/chapters/warnings. Real DOCX and EPUB files
# are built with python-docx and ebooklib -- the same libraries the app
# ships -- so the round trip is honest, not mocked.

import json

import pytest

from app.audiobook.extraction import extract_source, normalize_text


# ── Dispatch guards ───────────────────────────────────────────────────────────

def test_pdf_is_rejected_with_the_honest_deferred_message(tmp_path):
    pdf = tmp_path / "book.pdf"
    pdf.write_bytes(b"%PDF-1.4 fake")
    with pytest.raises(ValueError) as err:
        extract_source(str(pdf))
    assert "PDF import is not supported yet" in str(err.value)


def test_unknown_extension_is_rejected(tmp_path):
    weird = tmp_path / "book.xyz"
    weird.write_text("hello")
    with pytest.raises(ValueError) as err:
        extract_source(str(weird))
    assert "Unsupported manuscript format" in str(err.value)


def test_folder_without_project_json_is_rejected(tmp_path):
    with pytest.raises(ValueError) as err:
        extract_source(str(tmp_path))
    assert "not a Storythread project" in str(err.value)


# ── normalize_text ────────────────────────────────────────────────────────────

def test_normalize_collapses_blanks_and_line_endings():
    raw = "One\r\nTwo   \r\n\r\n\r\n\r\nThree\r"
    assert normalize_text(raw) == "One\nTwo\n\nThree"


def test_normalize_never_touches_punctuation():
    # Em dashes are the writer's own text at this layer -- house rule.
    raw = "She paused—then ran -- fast."
    assert normalize_text(raw) == raw


# ── TXT ───────────────────────────────────────────────────────────────────────

def test_txt_chapter_heading_lines_split_chapters(tmp_path):
    src = tmp_path / "my-novel.txt"
    src.write_text(
        "Dedication page text.\n\n"
        "Chapter 1\n\nFirst chapter prose.\n\n"
        "Chapter 2 -- The Fall\n\nSecond chapter prose.\n",
        encoding="utf-8",
    )
    result = extract_source(str(src))
    titles = [c.title for c in result.chapters]
    assert titles == ["Front Matter", "Chapter 1", "Chapter 2 -- The Fall"]
    assert result.chapters[1].text == "First chapter prose."
    assert result.title == "My Novel"          # prettified filename stem


def test_txt_without_headings_is_one_chapter_with_warning(tmp_path):
    src = tmp_path / "flat.txt"
    src.write_text("Just prose. No chapters anywhere.", encoding="utf-8")
    result = extract_source(str(src))
    assert len(result.chapters) == 1
    assert result.chapters[0].title == "Chapter 1"
    assert any("single chapter" in w for w in result.warnings)


def test_txt_sentence_starting_with_chapter_does_not_split(tmp_path):
    # A prose sentence beginning with "Chapter" must not be mistaken for a
    # heading -- headings are short whole lines.
    long_line = ("Chapter meetings bored her, and this sentence keeps rolling on "
                 "well past any plausible heading length so it cannot match the "
                 "heading pattern at all.")
    src = tmp_path / "one.txt"
    src.write_text(f"Chapter 1\n\n{long_line}\n\nChapter 2\n\nMore prose.\n", encoding="utf-8")
    result = extract_source(str(src))
    assert [c.title for c in result.chapters] == ["Chapter 1", "Chapter 2"]
    assert long_line in result.chapters[0].text


# ── Markdown ──────────────────────────────────────────────────────────────────

def test_markdown_h1_headings_become_chapters(tmp_path):
    src = tmp_path / "book.md"
    src.write_text(
        "# The Door\n\nProse one.\n\n# The Window\n\nProse two.\n",
        encoding="utf-8",
    )
    result = extract_source(str(src))
    assert [c.title for c in result.chapters] == ["The Door", "The Window"]


def test_markdown_falls_back_to_h2_when_no_h1(tmp_path):
    src = tmp_path / "book.md"
    src.write_text("## One\n\nA.\n\n## Two\n\nB.\n", encoding="utf-8")
    result = extract_source(str(src))
    assert [c.title for c in result.chapters] == ["One", "Two"]


def test_markdown_single_h1_is_promoted_to_book_title(tmp_path):
    src = tmp_path / "book.md"
    src.write_text("# The Hollow Road\n\nAll the prose.\n", encoding="utf-8")
    result = extract_source(str(src))
    assert result.title == "The Hollow Road"
    assert [c.title for c in result.chapters] == ["Chapter 1"]


# ── DOCX ──────────────────────────────────────────────────────────────────────

def test_docx_heading1_styles_split_chapters(tmp_path):
    from docx import Document
    doc = Document()
    doc.add_heading("Chapter One", level=1)
    doc.add_paragraph("First prose paragraph.")
    doc.add_paragraph("Second prose paragraph.")
    doc.add_heading("Chapter Two", level=1)
    doc.add_paragraph("More prose.")
    path = tmp_path / "book.docx"
    doc.save(str(path))

    result = extract_source(str(path))
    assert [c.title for c in result.chapters] == ["Chapter One", "Chapter Two"]
    assert "First prose paragraph." in result.chapters[0].text
    assert "More prose." in result.chapters[1].text


def test_docx_without_heading_styles_falls_back_to_heading_lines(tmp_path):
    from docx import Document
    doc = Document()
    doc.add_paragraph("Chapter 1")
    doc.add_paragraph("Prose in the first chapter.")
    doc.add_paragraph("Chapter 2")
    doc.add_paragraph("Prose in the second chapter.")
    path = tmp_path / "styled-less.docx"
    doc.save(str(path))

    result = extract_source(str(path))
    assert [c.title for c in result.chapters] == ["Chapter 1", "Chapter 2"]


# ── EPUB ──────────────────────────────────────────────────────────────────────

def _make_epub(tmp_path, chapters: list[tuple[str, str]]):
    from ebooklib import epub
    book = epub.EpubBook()
    book.set_identifier("test-epub")
    book.set_title("Ashes of Morning")
    book.set_language("en")
    book.add_author("Test Author")

    items = []
    for n, (title, body) in enumerate(chapters, start=1):
        item = epub.EpubHtml(title=title, file_name=f"ch{n}.xhtml", lang="en")
        item.content = f"<html><body><h1>{title}</h1><p>{body}</p></body></html>"
        book.add_item(item)
        items.append(item)

    book.toc = items
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())
    book.spine = ["nav"] + items
    path = tmp_path / "book.epub"
    epub.write_epub(str(path), book)
    return path


def test_epub_spine_documents_become_chapters(tmp_path):
    body = "The road disappeared beneath the gathering snow, mile after mile."
    path = _make_epub(tmp_path, [("The Door", body), ("The Window", body)])
    result = extract_source(str(path))
    assert result.title == "Ashes of Morning"
    assert result.author == "Test Author"
    assert [c.title for c in result.chapters] == ["The Door", "The Window"]
    # The heading is lifted into the title, not duplicated into the body.
    assert not result.chapters[0].text.startswith("The Door")
    assert "gathering snow" in result.chapters[0].text


def test_epub_furniture_pages_are_skipped(tmp_path):
    # The nav page ebooklib generates is short and must not become a chapter.
    body = "Long enough prose to clear the furniture threshold comfortably here."
    path = _make_epub(tmp_path, [("Only Chapter", body)])
    result = extract_source(str(path))
    assert [c.title for c in result.chapters] == ["Only Chapter"]


# ── Storythread project ───────────────────────────────────────────────────────

def _make_project(tmp_path, with_structure: bool):
    project = tmp_path / "MyNovel"
    (project / "manuscript").mkdir(parents=True)
    (project / "project.json").write_text(
        json.dumps({"title": "Curse of the Tomb Raider", "author": "Dean"}),
        encoding="utf-8",
    )
    (project / "manuscript" / "01-the-door.md").write_text(
        "# The Door\n\nProse of chapter one.\n", encoding="utf-8")
    (project / "manuscript" / "02-the-window.md").write_text(
        "Prose of chapter two, no heading line.\n", encoding="utf-8")
    if with_structure:
        # Reversed order proves the manifest, not the filenames, wins.
        (project / "manuscript" / "structure.json").write_text(json.dumps({
            "version": 1,
            "acts": [{"id": "a-1", "title": "Act I",
                      "chapters": ["02-the-window.md", "01-the-door.md"]}],
            "unassigned": [],
        }), encoding="utf-8")
    return project


def test_project_import_uses_structure_order(tmp_path):
    project = _make_project(tmp_path, with_structure=True)
    result = extract_source(str(project))
    assert result.title == "Curse of the Tomb Raider"
    assert result.author == "Dean"
    # structure.json order: window first, door second.
    assert [c.title for c in result.chapters] == ["The Window", "The Door"]
    # Heading lifted from the file that had one; stem-title for the other.
    assert "Prose of chapter one." in result.chapters[1].text


def test_project_import_without_structure_sorts_by_filename(tmp_path):
    project = _make_project(tmp_path, with_structure=False)
    result = extract_source(str(project))
    assert [c.title for c in result.chapters] == ["The Door", "The Window"]
