# routers/export.py -- Manuscript Export API
# ============================================
# This router handles exporting the writer's work out of the project
# into combined or snapshot files in the project's exports/ folder.
#
# Routes defined here:
#   POST /api/export/full-manuscript  -- combine all chapters into one file
#   POST /api/export/snapshot         -- save a dated copy of manuscript + metadata
#
# Full-manuscript supports four output formats (set via the `format` field):
#   markdown  (.md)  -- combine chapters separated by --- (original behavior)
#   txt       (.txt) -- same structure with Markdown syntax stripped
#   docx      (.docx)-- Word document with heading styles and inline formatting
#   epub      (.epub) -- EPUB e-book with chapters, TOC, and CSS
#
# Snapshot always produces a folder of .md files (it's a backup mechanism,
# not a publishing format).

import io
import json
import os
import re
import shutil
from datetime import datetime

from fastapi import APIRouter, HTTPException
from pydantic import BaseModel

from app.utils.structure_store import order_rank


# --- Router ---
router = APIRouter(prefix="/api/export", tags=["export"])


# --- Pydantic Models ---

class ExportRequest(BaseModel):
    """What the frontend sends when the writer clicks an export button.

    The `include_*` flags default to False so the old behavior (manuscript
    only) is preserved for any caller that doesn't pass the new fields.
    Callers that want the richer exports opt in explicitly.
    """
    folder_path: str   # The project's root directory (e.g. "C:/Users/.../MyNovel")

    # Output format for the full-manuscript export.
    # "markdown" is the original behavior. "txt", "docx", "epub" are new.
    # Extras (summaries, notes, profiles) are only appended in markdown/txt
    # exports -- DOCX and EPUB are prose-only clean publish formats.
    format: str = "markdown"   # "markdown" | "txt" | "docx" | "epub"

    # Opt-in extras for the richer Phase 6 export (markdown/txt only):
    #   chapter_summaries -> append a "Chapter Summaries" section to the full
    #                        manuscript, or copy summaries/chapters/ into the
    #                        snapshot folder.
    #   scene_summaries   -> same, but for summaries/scenes/<stem>/*.md
    #   notes             -> copy notes/*.md (outline, style guide, themes)
    #   profiles          -> copy profiles/ (characters, relationships, etc.)
    include_chapter_summaries: bool = False
    include_scene_summaries:   bool = False
    include_notes:             bool = False
    include_profiles:          bool = False

    # Optional chapter filter. When None or empty, ALL chapters in manuscript/
    # are exported (preserves the original behavior for callers that don't pass
    # the field). When non-empty, only chapters whose filename appears in this
    # list are included -- letting the writer export, say, just chapters 3-5
    # for sharing a draft excerpt without splitting their project.
    chapter_filenames: list[str] | None = None


class ExportResponse(BaseModel):
    """Confirmation returned after a successful export."""
    export_type: str    # "full-manuscript" or "snapshot"
    output_path: str    # Absolute path to the exported file or folder
    message: str        # Human-readable success message


# --- Helper Functions ---

def _exports_dir(folder_path: str) -> str:
    """
    Returns the absolute path to the exports/ folder inside a project.
    Raises 404 if the exports/ folder doesn't exist.
    """
    exports = os.path.join(folder_path, "exports")
    if not os.path.isdir(exports):
        raise HTTPException(
            status_code=404,
            detail=f"exports/ folder not found in: {folder_path}"
        )
    return exports


def _manuscript_dir(folder_path: str) -> str:
    """
    Returns the absolute path to the manuscript/ folder inside a project.
    Raises 404 if the manuscript/ folder doesn't exist.
    """
    manuscript = os.path.join(folder_path, "manuscript")
    if not os.path.isdir(manuscript):
        raise HTTPException(
            status_code=404,
            detail=f"manuscript/ folder not found in: {folder_path}"
        )
    return manuscript


def _collect_chapters(
    folder_path: str,
    only_filenames: list[str] | None = None,
) -> list[tuple[str, str]]:
    """
    Scans the manuscript/ folder for .md files, reads each one, and returns
    them as a sorted list of (filename, content) tuples.

    Sorting by filename ensures chapters come out in the right order when
    they're named like 01-chapter-one.md, 02-chapter-two.md, etc.

    When `only_filenames` is provided (and non-empty), only chapters whose
    filename appears in the set are returned. This powers the per-chapter
    selection UI in ExportModal -- the writer picks which chapters to export
    instead of the all-or-nothing behavior. None / empty list = include all
    (preserves the original behavior for older callers).
    """
    manuscript = _manuscript_dir(folder_path)
    chapters: list[tuple[str, str]] = []

    # Build a fast lookup set if a filter was supplied. Comparing filenames
    # is safe here -- the manuscript folder is flat (no subdirectories) so
    # there's no path-traversal risk from this check alone, and the
    # _manuscript_dir() guard already locks us inside the project.
    filter_set: set[str] | None = set(only_filenames) if only_filenames else None

    # os.scandir is efficient -- it reads directory entries without extra stat calls
    with os.scandir(manuscript) as entries:
        for entry in entries:
            if not (entry.is_file() and entry.name.endswith(".md")):
                continue
            if filter_set is not None and entry.name not in filter_set:
                continue
            try:
                with open(entry.path, "r", encoding="utf-8") as f:
                    content = f.read()
                chapters.append((entry.name, content))
            except OSError:
                # Skip files we can't read -- better to export what we can
                # than to fail the whole operation
                pass

    # Sort in reading order: the structure manifest is the ordering
    # authority when the project uses acts; filename order otherwise.
    # The exported book must match what the sidebar and Reader Mode show.
    rank = order_rank(folder_path)
    chapters.sort(key=lambda c: (rank.get(c[0], len(rank)), c[0]))
    return chapters


def _project_title(folder_path: str) -> str:
    """
    Reads the project title from project.json.
    Falls back to "untitled" if the file can't be read or has no title.
    """
    project_json_path = os.path.join(folder_path, "project.json")
    try:
        with open(project_json_path, "r", encoding="utf-8") as f:
            data = json.load(f)
        return data.get("title", "untitled") or "untitled"
    except (OSError, json.JSONDecodeError):
        return "untitled"


def _project_author(folder_path: str) -> str:
    """Read the author name from project.json, falling back to empty string."""
    project_json_path = os.path.join(folder_path, "project.json")
    try:
        with open(project_json_path, "r", encoding="utf-8") as f:
            data = json.load(f)
        return data.get("author", "") or ""
    except (OSError, json.JSONDecodeError):
        return ""


def _safe_title(title: str) -> str:
    """
    Converts a project title into a filename-safe slug.

    Examples:
      "The Lost Kingdom"    -> "the-lost-kingdom"
      "My Novel!!!"         -> "my-novel"
      "  Spaces & Symbols " -> "spaces-symbols"
    """
    # Lowercase, replace any run of non-alphanumeric characters with a single hyphen
    slug = re.sub(r"[^a-z0-9]+", "-", title.lower())
    # Strip leading/trailing hyphens
    return slug.strip("-") or "untitled"


# ── Phase 6: helpers for the opt-in extras ──────────────────────────────────
# These read optional project subfolders for the richer exports. Each returns
# an empty list/string when the folder is missing so the export never fails
# just because the writer hasn't created any summaries or notes yet.

def _collect_md_files(folder: str) -> list[tuple[str, str]]:
    """
    Generic "read all .md files in this folder" helper. Used for chapter
    summaries and notes. Sorted by filename so order is stable.
    Returns [] if the folder doesn't exist -- a missing subfolder is normal
    in projects where the writer hasn't used that feature yet.
    """
    if not os.path.isdir(folder):
        return []
    items: list[tuple[str, str]] = []
    with os.scandir(folder) as entries:
        for entry in entries:
            if entry.is_file() and entry.name.endswith(".md"):
                try:
                    with open(entry.path, "r", encoding="utf-8") as f:
                        items.append((entry.name, f.read()))
                except OSError:
                    pass
    items.sort(key=lambda c: c[0])
    return items


def _collect_scene_summaries(folder_path: str) -> list[tuple[str, int, str]]:
    """
    Walk summaries/scenes/<chapter-stem>/scene-NN.md and return a list of
    (chapter_stem, scene_index, content) tuples sorted by (chapter, index).
    Returns [] if no scene summaries exist yet.
    """
    scenes_root = os.path.join(folder_path, "summaries", "scenes")
    if not os.path.isdir(scenes_root):
        return []

    rows: list[tuple[str, int, str]] = []
    try:
        for chapter_stem in sorted(os.listdir(scenes_root)):
            chapter_dir = os.path.join(scenes_root, chapter_stem)
            if not os.path.isdir(chapter_dir):
                continue
            with os.scandir(chapter_dir) as entries:
                for entry in entries:
                    if not entry.is_file() or not entry.name.endswith(".md"):
                        continue
                    m = re.match(r"^scene-(\d{1,3})\.md$", entry.name)
                    if not m:
                        continue
                    try:
                        with open(entry.path, "r", encoding="utf-8") as f:
                            rows.append((chapter_stem, int(m.group(1)), f.read()))
                    except OSError:
                        pass
    except OSError:
        return []

    rows.sort(key=lambda r: (r[0], r[1]))
    return rows


def _copy_tree(src: str, dest: str) -> int:
    """
    Recursively copy src into dest. Returns the count of files copied.
    Returns 0 if src doesn't exist. Used by the snapshot export for
    summaries/, notes/, and profiles/.
    """
    if not os.path.isdir(src):
        return 0
    count = 0
    for root, _dirs, files in os.walk(src):
        rel = os.path.relpath(root, src)
        target_root = dest if rel == "." else os.path.join(dest, rel)
        os.makedirs(target_root, exist_ok=True)
        for name in files:
            try:
                shutil.copy2(os.path.join(root, name), os.path.join(target_root, name))
                count += 1
            except OSError:
                pass  # One bad file shouldn't abort the whole copy
    return count


# ── Format converters ────────────────────────────────────────────────────────
# Each converter takes the list of (filename, markdown_content) chapter tuples
# and produces either a string (for text-based formats) or bytes (for binary).


# Regex that matches the most common inline Markdown spans.
# Used by both the TXT stripper and the DOCX inline-formatter.
_INLINE_MD = re.compile(
    r'\*\*\*(?P<bold_italic>.+?)\*\*\*'
    r'|\*\*(?P<bold>.+?)\*\*'
    r'|\*(?P<italic>.+?)\*'
    r'|___(?P<u_bold_italic>.+?)___'
    r'|__(?P<u_bold>.+?)__'
    r'|_(?P<u_italic>.+?)_'
    r'|\[(?P<link_text>[^\]]+)\]\([^)]+\)'  # [text](url) -> text
    r'|`(?P<code>.+?)`',                    # `code` -> literal
    re.DOTALL,
)


def _strip_inline(text: str) -> str:
    """Remove inline Markdown syntax, keeping the visible text."""
    def repl(m: re.Match) -> str:
        for group in ("bold_italic", "bold", "italic",
                      "u_bold_italic", "u_bold", "u_italic",
                      "link_text", "code"):
            val = m.group(group)
            if val is not None:
                return val
        return m.group(0)
    return _INLINE_MD.sub(repl, text)


def _md_to_txt(markdown_text: str) -> str:
    """
    Convert a Markdown string to plain text by stripping all syntax.
    Preserves paragraph spacing and scene-break separators.
    """
    lines: list[str] = []
    for line in markdown_text.splitlines():
        stripped = line.rstrip()
        # Headings: remove the # markers, keep the title text
        m = re.match(r'^(#{1,6})\s+(.*)', stripped)
        if m:
            lines.append(m.group(2))
            continue
        # Scene breaks / horizontal rules -> blank separator line
        if re.match(r'^\s*(?:-{3,}|\*{3,})\s*$', stripped):
            lines.append("")
            continue
        # Blockquote markers
        stripped = re.sub(r'^>\s?', '', stripped)
        # Strip inline syntax
        stripped = _strip_inline(stripped)
        lines.append(stripped)
    return "\n".join(lines)


# ── DOCX builder ─────────────────────────────────────────────────────────────

def _docx_add_runs(paragraph, text: str) -> None:
    """
    Add bold / italic / plain runs to a python-docx paragraph based on
    the inline Markdown spans in `text`. Handles **, *, and combinations.
    Links become their link text; code spans stay as plain text.
    """
    pos = 0
    for m in _INLINE_MD.finditer(text):
        # Add any plain text before this span
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])

        # Determine which group matched and extract the visible text
        if m.group("bold_italic") or m.group("u_bold_italic"):
            run = paragraph.add_run(m.group("bold_italic") or m.group("u_bold_italic"))
            run.bold = True
            run.italic = True
        elif m.group("bold") or m.group("u_bold"):
            run = paragraph.add_run(m.group("bold") or m.group("u_bold"))
            run.bold = True
        elif m.group("italic") or m.group("u_italic"):
            run = paragraph.add_run(m.group("italic") or m.group("u_italic"))
            run.italic = True
        elif m.group("link_text"):
            paragraph.add_run(m.group("link_text"))
        elif m.group("code"):
            paragraph.add_run(m.group("code"))

        pos = m.end()

    # Add any trailing plain text after the last span
    if pos < len(text):
        paragraph.add_run(text[pos:])


def _build_docx(chapters: list[tuple[str, str]], title: str) -> bytes:
    """
    Build a Word (.docx) document from the chapter list.

    Structure:
      - Document title as a Title-style paragraph at the very top
      - Each chapter starts on a new page (except the first)
      - Chapter headings map to Word Heading 1 / 2 / 3 styles
      - Scene breaks (---) become centered "* * *" dividers
      - Body lines get inline bold/italic from their Markdown spans
    """
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Document title
    title_para = doc.add_paragraph(title, style="Title")
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for i, (_filename, content) in enumerate(chapters):
        if i > 0:
            doc.add_page_break()

        blank_run = False   # track consecutive blanks to avoid extra spacing

        for raw_line in content.splitlines():
            line = raw_line.rstrip()

            # Heading 1: chapter title (# Title)
            m = re.match(r'^(#{1,3})\s+(.*)', line)
            if m:
                level = len(m.group(1))
                doc.add_heading(m.group(2), level=level)
                blank_run = False
                continue

            # Scene break: --- or *** -> centered divider
            if re.match(r'^\s*(?:-{3,}|\*{3,})\s*$', line):
                p = doc.add_paragraph("* * *")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                blank_run = False
                continue

            # Blank line: one blank paragraph max, skip consecutive blanks
            if not line.strip():
                if not blank_run:
                    doc.add_paragraph("")
                    blank_run = True
                continue

            blank_run = False

            # Body paragraph: add inline formatted runs. _docx_add_runs handles
            # both the styled spans and any plain text between them.
            p = doc.add_paragraph()
            p.style = "Normal"
            _docx_add_runs(p, line)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── EPUB builder ─────────────────────────────────────────────────────────────

# Minimal CSS injected into every EPUB chapter. Keeps the prose readable
# across e-reader apps without fighting their default stylesheets.
_EPUB_CSS = """\
body {
  font-family: Georgia, "Times New Roman", serif;
  font-size: 1em;
  line-height: 1.6;
  margin: 0 1em;
}
h1 { font-size: 1.4em; margin: 2em 0 0.5em; }
h2 { font-size: 1.2em; margin: 1.5em 0 0.4em; }
h3 { font-size: 1.05em; margin: 1.2em 0 0.3em; }
p  { margin: 0; text-indent: 1.5em; }
p.scene-break { text-align: center; text-indent: 0; margin: 1em 0; }
"""


def _extract_first_heading(md_text: str) -> str | None:
    """Return the text of the first # heading in a Markdown string, or None."""
    for line in md_text.splitlines():
        m = re.match(r'^#{1,3}\s+(.*)', line.rstrip())
        if m:
            return m.group(1).strip()
    return None


def _build_epub(chapters: list[tuple[str, str]], title: str, author: str = "") -> bytes:
    """
    Build an EPUB e-book from the chapter list using ebooklib.

    Each chapter file becomes one XHTML chapter in the EPUB spine.
    Markdown is converted to HTML via the `markdown` library so formatting
    (bold, italic, headings, scene breaks) renders correctly in e-readers.
    """
    import markdown as md_lib
    from ebooklib import epub

    book = epub.EpubBook()
    book.set_title(title)
    book.set_language("en")
    if author:
        book.add_author(author)

    # Global stylesheet
    css = epub.EpubItem(
        uid="style",
        file_name="style/main.css",
        media_type="text/css",
        content=_EPUB_CSS,
    )
    book.add_item(css)

    spine_items: list = ["nav"]
    toc_items:   list = []

    md_converter = md_lib.Markdown(extensions=["extra"])

    for i, (filename, content) in enumerate(chapters):
        chapter_num   = i + 1
        chapter_title = _extract_first_heading(content) or f"Chapter {chapter_num}"

        # Convert Markdown -> HTML. The `extra` extension handles tables,
        # fenced code blocks, attribute lists, etc. Most fiction won't use
        # these, but supporting them costs nothing.
        md_converter.reset()
        body_html = md_converter.convert(content)

        # Replace bare <hr> elements (from --- scene breaks) with a
        # styled paragraph so e-readers that ignore <hr> CSS still show them
        body_html = body_html.replace("<hr />", '<p class="scene-break">* * *</p>')
        body_html = body_html.replace("<hr>",   '<p class="scene-break">* * *</p>')

        xhtml = (
            '<?xml version="1.0" encoding="utf-8"?>\n'
            '<!DOCTYPE html>\n'
            '<html xmlns="http://www.w3.org/1999/xhtml">\n'
            '<head>\n'
            f'  <title>{chapter_title}</title>\n'
            '  <link rel="stylesheet" href="../style/main.css" type="text/css"/>\n'
            '</head>\n'
            '<body>\n'
            f'{body_html}\n'
            '</body>\n'
            '</html>'
        )

        epub_chapter = epub.EpubHtml(
            title     = chapter_title,
            file_name = f"chapter_{chapter_num:03d}.xhtml",
            lang      = "en",
        )
        # ebooklib 0.18+ expects content as bytes, not str
        epub_chapter.content = xhtml.encode("utf-8")
        epub_chapter.add_item(css)
        book.add_item(epub_chapter)
        spine_items.append(epub_chapter)
        toc_items.append(epub.Link(
            f"chapter_{chapter_num:03d}.xhtml",
            chapter_title,
            f"chapter_{chapter_num:03d}",
        ))

    book.toc   = toc_items
    book.spine = spine_items
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())

    buf = io.BytesIO()
    epub.write_epub(buf, book)
    return buf.getvalue()


# --- POST /api/export/full-manuscript ---

@router.post("/full-manuscript", response_model=ExportResponse)
async def export_full_manuscript(request: ExportRequest):
    """
    Combines all chapters from manuscript/ into a single file.

    Four output formats are supported (request.format):
      markdown (.md)  -- chapters joined by --- separators (original behavior)
      txt      (.txt) -- same structure, Markdown syntax stripped to plain text
      docx     (.docx)-- Word document with heading styles and inline formatting
      epub     (.epub)-- EPUB e-book with chapters, TOC, and CSS

    For markdown and txt, opt-in extras (summaries, notes, profiles) can be
    appended as sections. For docx and epub the export is prose-only -- extras
    are working documents that don't belong in a clean publish file.

    The output is always written to exports/{slug}-full-manuscript.{ext} and
    is overwritten on each export, giving the writer one canonical publish file.
    """
    fmt = (request.format or "markdown").lower()
    if fmt not in ("markdown", "txt", "docx", "epub"):
        raise HTTPException(status_code=400, detail=f"Unknown format '{fmt}'. Use markdown, txt, docx, or epub.")

    folder_path = request.folder_path
    exports = _exports_dir(folder_path)

    chapters = _collect_chapters(folder_path, request.chapter_filenames)
    if not chapters:
        if request.chapter_filenames:
            raise HTTPException(
                status_code=404,
                detail="None of the selected chapter filenames matched files in manuscript/."
            )
        raise HTTPException(
            status_code=404,
            detail="No chapters found in manuscript/ folder. Write some chapters first!"
        )

    title = _project_title(folder_path)
    slug  = _safe_title(title)

    # ── DOCX ──────────────────────────────────────────────────────────────────
    if fmt == "docx":
        try:
            data = _build_docx(chapters, title)
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"DOCX build failed: {e}")

        output_filename = f"{slug}-full-manuscript.docx"
        output_path     = os.path.join(exports, output_filename)
        try:
            with open(output_path, "wb") as f:
                f.write(data)
        except OSError as e:
            raise HTTPException(status_code=500, detail=f"Could not write export file: {e}")

        return ExportResponse(
            export_type="full-manuscript",
            output_path=output_path,
            message=f"Exported {len(chapters)} chapter(s) to {output_filename}",
        )

    # ── EPUB ──────────────────────────────────────────────────────────────────
    if fmt == "epub":
        try:
            author = _project_author(folder_path)
            data   = _build_epub(chapters, title, author)
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"EPUB build failed: {e}")

        output_filename = f"{slug}-full-manuscript.epub"
        output_path     = os.path.join(exports, output_filename)
        try:
            with open(output_path, "wb") as f:
                f.write(data)
        except OSError as e:
            raise HTTPException(status_code=500, detail=f"Could not write export file: {e}")

        return ExportResponse(
            export_type="full-manuscript",
            output_path=output_path,
            message=f"Exported {len(chapters)} chapter(s) to {output_filename}",
        )

    # ── Markdown / TXT (shared path) ──────────────────────────────────────────
    # Build the combined Markdown text first, then optionally strip it for TXT.
    parts: list[str] = []
    for _filename, content in chapters:
        parts.append(content.strip())
    combined_md = "\n\n---\n\n".join(parts) + "\n"

    # Opt-in appendices (markdown and txt only)
    appendices: list[str] = []
    extras_summary_parts: list[str] = []

    if request.include_chapter_summaries:
        chapter_summaries = _collect_md_files(os.path.join(folder_path, "summaries", "chapters"))
        if chapter_summaries:
            section = ["# Chapter Summaries"]
            for name, body in chapter_summaries:
                section.append(f"\n## {name.removesuffix('.md')}\n")
                section.append(body.strip())
            appendices.append("\n".join(section))
            extras_summary_parts.append(f"{len(chapter_summaries)} chapter summaries")

    if request.include_scene_summaries:
        scene_rows = _collect_scene_summaries(folder_path)
        if scene_rows:
            section = ["# Scene Summaries"]
            current_stem: str | None = None
            for stem, idx, body in scene_rows:
                if stem != current_stem:
                    section.append(f"\n## {stem}\n")
                    current_stem = stem
                section.append(f"\n### Scene {idx}\n")
                trimmed = re.sub(r"^\s*#[^\n]*\n", "", body.strip(), count=1)
                section.append(trimmed.strip())
            appendices.append("\n".join(section))
            extras_summary_parts.append(f"{len(scene_rows)} scene summaries")

    if request.include_notes:
        notes = _collect_md_files(os.path.join(folder_path, "notes"))
        if notes:
            section = ["# Notes"]
            for name, body in notes:
                section.append(f"\n## {name.removesuffix('.md')}\n")
                section.append(body.strip())
            appendices.append("\n".join(section))
            extras_summary_parts.append(f"{len(notes)} notes")

    if request.include_profiles:
        profile_types = ["character", "relationship", "location", "lore"]
        profile_chunks: list[str] = []
        total_profiles = 0
        for ptype in profile_types:
            files = _collect_md_files(os.path.join(folder_path, "profiles", ptype))
            if not files:
                continue
            profile_chunks.append(f"\n## {ptype.title()}s\n")
            for name, body in files:
                profile_chunks.append(f"\n### {name.removesuffix('.md')}\n")
                profile_chunks.append(body.strip())
                total_profiles += 1
        if profile_chunks:
            appendices.append("# Profiles" + "".join(profile_chunks))
            extras_summary_parts.append(f"{total_profiles} profiles")

    if appendices:
        combined_md = combined_md.rstrip() + "\n\n---\n\n" + "\n\n---\n\n".join(appendices) + "\n"

    # Choose extension and apply text transformation
    if fmt == "txt":
        output_text     = _md_to_txt(combined_md)
        output_filename = f"{slug}-full-manuscript.txt"
        encoding        = "utf-8"
    else:
        output_text     = combined_md
        output_filename = f"{slug}-full-manuscript.md"
        encoding        = "utf-8"

    output_path = os.path.join(exports, output_filename)
    try:
        with open(output_path, "w", encoding=encoding) as f:
            f.write(output_text)
    except OSError as e:
        raise HTTPException(status_code=500, detail=f"Could not write export file: {e}")

    extras_msg = ""
    if extras_summary_parts:
        extras_msg = " (plus " + ", ".join(extras_summary_parts) + ")"

    return ExportResponse(
        export_type="full-manuscript",
        output_path=output_path,
        message=f"Exported {len(chapters)} chapter(s) to {output_filename}{extras_msg}",
    )


# --- POST /api/export/snapshot ---

@router.post("/snapshot", response_model=ExportResponse)
async def export_snapshot(request: ExportRequest):
    """
    Creates a dated snapshot of the manuscript and project metadata.

    Unlike full-manuscript export (which overwrites one file), snapshots
    accumulate over time. Each snapshot gets its own timestamped folder
    inside exports/, so the writer can look back at earlier versions.

    The snapshot includes:
      - All .md files from manuscript/ (chapter drafts)
      - project.json (project settings and metadata)
    """
    folder_path = request.folder_path
    exports = _exports_dir(folder_path)
    manuscript = _manuscript_dir(folder_path)

    # Generate a timestamped folder name
    title = _project_title(folder_path)
    slug = _safe_title(title)
    timestamp = datetime.now().strftime("%Y-%m-%d-%H%M%S")
    snapshot_name = f"{slug}-snapshot-{timestamp}"
    snapshot_dir = os.path.join(exports, snapshot_name)

    # Create the snapshot folder
    try:
        os.makedirs(snapshot_dir, exist_ok=True)
    except OSError as e:
        raise HTTPException(
            status_code=500,
            detail=f"Could not create snapshot folder: {e}"
        )

    copied_count = 0

    # Optional per-chapter filter: when the writer chose specific chapters in
    # ExportModal, the snapshot only mirrors those (the other manuscript files
    # are simply not copied). None / empty list = full snapshot, same as before.
    chapter_filter: set[str] | None = (
        set(request.chapter_filenames) if request.chapter_filenames else None
    )

    # Copy selected .md files from manuscript/ into the snapshot
    try:
        with os.scandir(manuscript) as entries:
            for entry in entries:
                if not (entry.is_file() and entry.name.endswith(".md")):
                    continue
                if chapter_filter is not None and entry.name not in chapter_filter:
                    continue
                shutil.copy2(entry.path, os.path.join(snapshot_dir, entry.name))
                copied_count += 1
    except OSError as e:
        raise HTTPException(
            status_code=500,
            detail=f"Error copying manuscript files: {e}"
        )

    # Copy project.json if it exists
    project_json = os.path.join(folder_path, "project.json")
    if os.path.isfile(project_json):
        try:
            shutil.copy2(project_json, os.path.join(snapshot_dir, "project.json"))
        except OSError:
            pass  # Not critical -- the chapters are what matter most

    # ── Opt-in extras (Phase 6). Each bucket is copied into a matching
    #    subfolder inside the snapshot dir so the shape mirrors the project
    #    layout -- restoring from a snapshot is then a straight copy-paste.
    extras_parts: list[str] = []

    if request.include_chapter_summaries:
        n = _copy_tree(
            os.path.join(folder_path, "summaries", "chapters"),
            os.path.join(snapshot_dir, "summaries", "chapters"),
        )
        if n:
            extras_parts.append(f"{n} chapter summaries")

    if request.include_scene_summaries:
        n = _copy_tree(
            os.path.join(folder_path, "summaries", "scenes"),
            os.path.join(snapshot_dir, "summaries", "scenes"),
        )
        if n:
            extras_parts.append(f"{n} scene summary files")

    if request.include_notes:
        n = _copy_tree(
            os.path.join(folder_path, "notes"),
            os.path.join(snapshot_dir, "notes"),
        )
        if n:
            extras_parts.append(f"{n} notes")

    if request.include_profiles:
        n = _copy_tree(
            os.path.join(folder_path, "profiles"),
            os.path.join(snapshot_dir, "profiles"),
        )
        if n:
            extras_parts.append(f"{n} profile files")

    if copied_count == 0:
        # Clean up the empty folder since there was nothing to snapshot
        try:
            shutil.rmtree(snapshot_dir)
        except OSError:
            pass
        if chapter_filter is not None:
            raise HTTPException(
                status_code=404,
                detail="None of the selected chapter filenames matched files in manuscript/."
            )
        raise HTTPException(
            status_code=404,
            detail="No chapters found in manuscript/ folder. Nothing to snapshot."
        )

    extras_msg = ""
    if extras_parts:
        extras_msg = " (plus " + ", ".join(extras_parts) + ")"

    return ExportResponse(
        export_type="snapshot",
        output_path=snapshot_dir,
        message=f"Snapshot saved: {copied_count} chapter(s) + project.json to {snapshot_name}/{extras_msg}",
    )
