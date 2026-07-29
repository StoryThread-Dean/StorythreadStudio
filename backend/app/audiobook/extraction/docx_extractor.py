# extraction/docx_extractor.py -- Word manuscripts via python-docx.
# ==================================================================
# python-docx is already a project dependency (the export feature writes
# DOCX with it); importing is the same library driven in reverse. Chapter
# detection uses Word's own structure first: paragraphs styled "Heading 1"
# are chapter titles. Manuscripts written without heading styles fall back
# to the plain-text detector's idea of a heading line.

import re

from docx import Document

from app.audiobook.extraction import ExtractedChapter, ExtractionResult, normalize_text
from app.audiobook.extraction.text_extractor import _TXT_HEADING_RE, _pretty_stem


def extract_docx(path: str) -> ExtractionResult:
    doc = Document(path)

    # Word files carry metadata ("core properties") -- often blank, but when
    # an author filled in File > Info we get title/author for free.
    props = doc.core_properties
    result = ExtractionResult(
        title=(props.title or "").strip() or _pretty_stem(path),
        author=(props.author or "").strip(),
    )

    # Walk paragraphs once, splitting on Heading 1 styles. style can be None
    # for exotic documents, so guard every attribute step.
    chapters: list[ExtractedChapter] = []
    current_title: str | None = None
    current_lines: list[str] = []
    saw_heading_style = False

    def flush() -> None:
        body = normalize_text("\n\n".join(current_lines))
        if current_title is None and not body:
            return                       # nothing before the first heading
        chapters.append(ExtractedChapter(
            title=current_title if current_title is not None else "Front Matter",
            text=body,
        ))

    for para in doc.paragraphs:
        style_name = getattr(getattr(para, "style", None), "name", "") or ""
        text = para.text.strip()
        if style_name.startswith("Heading 1"):
            saw_heading_style = True
            flush()
            current_title = text or "Untitled Chapter"
            current_lines = []
        elif text:
            current_lines.append(text)
    flush()

    if saw_heading_style:
        result.chapters = chapters
        return result

    # No Heading 1 styles anywhere: re-detect using plain-text heading lines
    # over the flattened text, exactly like a .txt import.
    flat = normalize_text("\n\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip()))
    lines = flat.split("\n")
    heading_idxs = [i for i, line in enumerate(lines) if _TXT_HEADING_RE.match(line)]

    if len(heading_idxs) < 2:
        result.chapters = [ExtractedChapter(title="Chapter 1", text=flat)]
        result.warnings.append(
            "No Heading 1 styles or chapter heading lines were detected; the "
            "manuscript was imported as a single chapter."
        )
        return result

    preamble = "\n".join(lines[: heading_idxs[0]]).strip("\n")
    if preamble:
        result.chapters.append(ExtractedChapter(title="Front Matter", text=preamble))
    for n, start in enumerate(heading_idxs):
        end = heading_idxs[n + 1] if n + 1 < len(heading_idxs) else len(lines)
        result.chapters.append(ExtractedChapter(
            title=lines[start].strip(),
            text="\n".join(lines[start + 1 : end]).strip("\n"),
        ))
    return result


# Re-exported for tests that want to sanity-check the fallback regex is the
# shared one (a single source of truth for "what looks like a heading").
HEADING_LINE_RE: re.Pattern[str] = _TXT_HEADING_RE
